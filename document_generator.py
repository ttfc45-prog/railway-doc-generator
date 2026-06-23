"""
document_generator.py
Railway Documentation Generator - Document Assembly Engine
Builds .docx (Word) and PDF documents from templates, LLM text, and computed tables/figures.
"""

import io
import datetime
from pathlib import Path
from typing import Callable

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from config import (
    OUTPUT_DOCX, OUTPUT_PDF, FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_H1,
    FONT_SIZE_H2, FONT_SIZE_H3, FONT_SIZE_CAPTION, HEADER_COLOR,
    TABLE_HEADER_BG, TABLE_ALT_BG, DOCUMENT_TYPES, FIGURE_DPI
)
from project_database import ProjectDatabase as PDB
from calculations import CalculationEngine, RailwayCalculations
from project_model import ProjectModel
from llm_writer import build_narrative_context
from templates import TEMPLATES
from tables import TableGenerator
from figures import generate_all_figures
import llm_writer


# ═══════════════════════════════════════════════════════════════
# HELPER – RGB from hex
# ═══════════════════════════════════════════════════════════════

def _rgb(hex6: str) -> RGBColor:
    r = int(hex6[0:2], 16)
    g = int(hex6[2:4], 16)
    b = int(hex6[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_bg(cell, hex6: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6)
    tcPr.append(shd)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════

class DocumentGenerator:
    """
    Orchestrates the assembly of a professional Word document from the
    template definition, computed data, LLM-written chapters, tables, and figures.
    """

    def __init__(self, doc_key: str, p: dict | None = None):
        self.doc_key    = doc_key
        self.doc_title  = DOCUMENT_TYPES.get(doc_key, doc_key)
        raw             = p or PDB.get_all()
        # Strip forbidden keys before building ProjectModel
        self.model      = ProjectModel({k: v for k, v in raw.items()
                                        if k not in ProjectModel._FORBIDDEN_KEYS})
        self.cs         = CalculationEngine.run(self.model)
        self.ctx        = build_narrative_context(self.model, self.cs)
        # Convenience aliases (kept for Word styling code that references self.p for metadata)
        self.p          = self.model.to_dict()
        self.ops        = self.cs.ops
        self.ram        = self.cs.ram
        self.headway    = self.cs.headway
        self.traction   = self.cs.traction
        self.fig_counter= 0
        self.tab_counter= 0
        self.doc: Document | None = None

    # ───────────────────────────────────────────
    # Core entry point
    # ───────────────────────────────────────────

    def build(self, use_llm: bool = True, progress_cb: Callable | None = None) -> Path:
        """
        Build the Word document and return its file path.

        Parameters
        ----------
        use_llm      : whether to call the LLM for chapter text (slow)
        progress_cb  : optional callable(pct: float, msg: str) for UI progress
        """
        self.fig_counter = 0
        self.tab_counter = 0
        self.doc = Document()
        self._configure_page()
        self._add_cover_page()
        self._add_document_info_table()
        self._add_page_break()
        self._add_toc_placeholder()
        self._add_page_break()

        template = TEMPLATES.get(self.doc_key, TEMPLATES["BOD"])
        n = len(template)

        # Pre-generate figures
        figs = generate_all_figures(self.p, self.ops, self.ram, self.headway)
        # Merge specialist figures (Phase 1.5 six new document types)
        try:
            from figures import generate_specialist_figures
            figs.update(generate_specialist_figures(self.cs))
        except Exception:
            pass

        for idx, section in enumerate(template):
            if progress_cb:
                progress_cb((idx + 1) / n, f"Writing: {section['title']}")
            self._write_section(section, figs, use_llm)

        path = OUTPUT_DOCX / f"{self.p.get('project_name','Project').replace(' ','_')}_{self.doc_key}.docx"
        self.doc.save(str(path))
        return path

    # ───────────────────────────────────────────
    # Page / style setup
    # ───────────────────────────────────────────

    def _configure_page(self):
        section = self.doc.sections[0]
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin   = Cm(2.5)
        section.bottom_margin= Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)
        self._add_header_footer()

    def _add_header_footer(self):
        hdr = self.doc.sections[0].header
        p   = hdr.paragraphs[0]
        p.clear()
        run = p.add_run(
            f"{self.p.get('project_name','')} | {self.doc_title} | {self.p.get('document_number','')}"
        )
        run.font.name  = FONT_NAME
        run.font.size  = Pt(8)
        run.font.color.rgb = _rgb("888888")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.border_bottom = True

        ftr = self.doc.sections[0].footer
        fp  = ftr.paragraphs[0]
        fp.clear()
        fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = fp.add_run(
            f"© {self.p.get('consultant','Railway Consultants')} | "
            f"Rev {self.p.get('revision','A')} | {self.p.get('status','Draft')} | "
            f"Page "
        )
        run2.font.name = FONT_NAME
        run2.font.size = Pt(8)
        run2.font.color.rgb = _rgb("888888")
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        fp.runs[-1]._r.append(fld)
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fp.runs[-1]._r.append(instrText)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        fp.runs[-1]._r.append(fld2)

    # ───────────────────────────────────────────
    # Cover Page
    # ───────────────────────────────────────────

    def _add_cover_page(self):
        doc = self.doc
        doc.add_paragraph()
        doc.add_paragraph()

        # Company name
        comp = doc.add_paragraph()
        comp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = comp.add_run(self.p.get("consultant", "Railway Engineering Consultants").upper())
        r.font.name  = FONT_NAME
        r.font.size  = Pt(14)
        r.font.bold  = True
        r.font.color.rgb = _rgb(HEADER_COLOR)

        doc.add_paragraph()

        # Document title
        title_p = doc.add_paragraph()
        title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = title_p.add_run(self.doc_title.upper())
        rt.font.name  = FONT_NAME
        rt.font.size  = Pt(20)
        rt.font.bold  = True
        rt.font.color.rgb = _rgb(HEADER_COLOR)

        # Subtitle
        sub = doc.add_paragraph()
        sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = sub.add_run(
            f"{self.p.get('project_name','')} — {self.p.get('line_name','')} | {self.p.get('country','')}"
        )
        rs.font.name  = FONT_NAME
        rs.font.size  = Pt(13)
        rs.font.color.rgb = _rgb("444444")

        doc.add_paragraph()
        doc.add_paragraph()

        # Info box
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        info = [
            ("Client",          self.p.get("client", "")),
            ("Document No.",    self.p.get("document_number", "")),
            ("Revision",        self.p.get("revision", "A")),
            ("Status",          self.p.get("status", "Draft")),
            ("Date",            datetime.date.today().strftime("%d %B %Y")),
            ("Prepared by",     self.p.get("consultant", "")),
        ]
        for i, (label, val) in enumerate(info):
            lc = table.cell(i, 0)
            vc = table.cell(i, 1)
            _set_cell_bg(lc, "EEF2FB")
            lc.text = label
            vc.text = val
            for cell in [lc, vc]:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)

    # ───────────────────────────────────────────
    # Document info table (revision history)
    # ───────────────────────────────────────────

    def _add_document_info_table(self):
        doc = self.doc
        doc.add_paragraph()
        h = doc.add_paragraph("Revision History")
        h.runs[0].font.bold = True
        h.runs[0].font.size = Pt(11)

        table = doc.add_table(rows=3, cols=5)
        table.style = "Table Grid"
        headers = ["Rev", "Date", "Description", "Author", "Approved"]
        for j, hdr in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = hdr
            _set_cell_bg(cell, TABLE_HEADER_BG)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
        row_data = [
            ["A", datetime.date.today().strftime("%d/%m/%Y"), "First issue for review", "", ""],
            ["B", "", "Issued for approval", "", ""],
        ]
        for i, rd in enumerate(row_data):
            for j, val in enumerate(rd):
                table.cell(i+1, j).text = val

    # ───────────────────────────────────────────
    # TOC placeholder
    # ───────────────────────────────────────────

    def _add_toc_placeholder(self):
        doc = self.doc
        h = doc.add_paragraph("Table of Contents")
        h.runs[0].font.bold = True
        h.runs[0].font.size = Pt(14)
        doc.add_paragraph("[ This table of contents is generated automatically by Microsoft Word. "
                          "Right-click and select 'Update Field' to refresh. ]")

    # ───────────────────────────────────────────
    # Section writer
    # ───────────────────────────────────────────

    def _write_section(self, section: dict, figs: dict, use_llm: bool):
        doc   = self.doc
        level = section.get("level", 1)
        title = section["title"]

        # Heading
        style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 1")
        h = doc.add_heading(title, level=level)
        for run in h.runs:
            run.font.name  = FONT_NAME
            run.font.color.rgb = _rgb(HEADER_COLOR)

        # Body text from generator
        gen_key = section.get("generator")
        if gen_key and use_llm:
            try:
                gen_fn = llm_writer.CHAPTER_GENERATORS.get(gen_key)
                if gen_fn:
                    if gen_key in ("introduction", "scope", "conclusion"):
                        text = gen_fn(self.ctx, self.doc_title)
                    else:
                        text = gen_fn(self.ctx)
                else:
                    text = ""
            except Exception:
                text = ""

            if text:
                for para_text in text.split("\n\n"):
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.paragraph_format.space_after = Pt(6)
                        for run in p.runs:
                            run.font.name = FONT_NAME
                            run.font.size = Pt(FONT_SIZE_BODY)

        # Tables
        for table_key in section.get("tables", []):
            df = self._get_table(table_key)
            if df is not None and not df.empty:
                self._add_dataframe_table(df, table_key)

        # Figures
        for fig_key in section.get("figures", []):
            fig_path = figs.get(fig_key)
            if fig_path and Path(fig_path).exists():
                self._add_figure(fig_path, fig_key)

    # ───────────────────────────────────────────
    # Table router
    # ───────────────────────────────────────────

    def _get_table(self, key: str) -> pd.DataFrame | None:
        tg = TableGenerator()
        dispatch = {
            "project_data_table":            lambda: tg.project_data_table(self.model, self.cs),
            "operational_parameters_table":  lambda: tg.operational_parameters_table(self.model, self.cs),
            "rolling_stock_table":           lambda: tg.rolling_stock_table(self.model, self.cs),
            "fleet_calculation_table":       lambda: tg.fleet_calculation_table(self.model, self.cs),
            "headway_breakdown_table":       lambda: tg.headway_breakdown_table(self.cs),
            "ram_targets_table":             lambda: tg.ram_targets_table(self.model, self.cs),
            "subsystem_availability_table":  lambda: tg.subsystem_availability_table(self.cs),
            "fmeca_table":                   lambda: tg.fmeca_table(),
            "hazard_log_table":              lambda: tg.hazard_log_table(self.cs),
            "interface_matrix_table":        lambda: tg.interface_matrix_table(),
            "srs_requirements_table":        lambda: tg.srs_requirements_table(self.model, self.cs),
            "standards_table":               lambda: tg.standards_table(self.model),
            "environmental_conditions_table":lambda: tg.environmental_conditions_table(self.model),
            "performance_kpi_table":         lambda: tg.performance_kpi_table(self.cs),
            "traction_parameters_table":     lambda: tg.traction_parameters_table(self.model, self.cs),
            "capacity_study_table":          lambda: tg.capacity_study_table(self.model, self.cs),
            # Reliability Report
            "reliability_mtbf_allocation_table":  lambda: tg.reliability_mtbf_allocation_table(self.cs),
            "reliability_R_t_table":              lambda: tg.reliability_R_t_table(self.cs),
            "reliability_block_diagram_table":    lambda: tg.reliability_block_diagram_table(self.cs),
            # Availability Report
            "availability_budget_table":          lambda: tg.availability_budget_table(self.model, self.cs),
            "availability_waterfall_table":       lambda: tg.availability_waterfall_table(self.cs),
            "availability_subsystem_comparison_table": lambda: tg.availability_subsystem_comparison_table(self.cs),
            # Maintainability Study
            "maintainability_M_t_table":          lambda: tg.maintainability_M_t_table(self.cs),
            "maintenance_levels_table":           lambda: tg.maintenance_levels_table(self.model, self.cs),
            "corrective_vs_preventive_table":     lambda: tg.corrective_vs_preventive_table(self.cs),
            # Hazard Log (with CS links)
            "hazard_log_with_cs_table":           lambda: tg.hazard_log_with_cs_table(self.cs),
            # Energy Management Plan
            "energy_balance_table":               lambda: tg.energy_balance_table(self.model, self.cs),
            "substation_sizing_table":            lambda: tg.substation_sizing_table(self.model, self.cs),
            "energy_kpi_table":                   lambda: tg.energy_kpi_table(self.cs),
            # Human Factors Report
            "occ_workload_table":                 lambda: tg.occ_workload_table(self.model, self.cs),
            "staffing_table":                     lambda: tg.staffing_table(self.model, self.cs),
            "evacuation_assumptions_table":       lambda: tg.evacuation_assumptions_table(self.model, self.cs),
            "operator_responsibilities_table":    lambda: tg.operator_responsibilities_table(),
        }
        fn = dispatch.get(key)
        return fn() if fn else None

    # ───────────────────────────────────────────
    # DataFrame → Word table
    # ───────────────────────────────────────────

    def _add_dataframe_table(self, df: pd.DataFrame, key: str):
        doc = self.doc
        self.tab_counter += 1
        caption_txt = key.replace("_table","").replace("_"," ").title()

        # Caption above
        cap = doc.add_paragraph(f"Table {self.tab_counter}: {caption_txt}")
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in cap.runs:
            run.font.name  = FONT_NAME
            run.font.size  = Pt(FONT_SIZE_CAPTION)
            run.font.italic = True
            run.font.bold   = False

        rows, cols = df.shape
        table = doc.add_table(rows=rows + 1, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for j, col_name in enumerate(df.columns):
            cell = table.cell(0, j)
            cell.text = str(col_name)
            _set_cell_bg(cell, TABLE_HEADER_BG)
            for para in cell.paragraphs:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name  = FONT_NAME
                    run.font.bold  = True
                    run.font.size  = Pt(8)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # Data rows
        for i, row in df.iterrows():
            bg = TABLE_ALT_BG if i % 2 == 0 else "FFFFFF"
            for j, val in enumerate(row):
                cell = table.cell(i + 1, j)
                cell.text = str(val) if val is not None else ""
                _set_cell_bg(cell, bg)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(8)

        doc.add_paragraph()  # spacing after table

    # ───────────────────────────────────────────
    # Figure inserter
    # ───────────────────────────────────────────

    def _add_figure(self, fig_path: Path, key: str):
        doc = self.doc
        self.fig_counter += 1
        caption_txt = key.replace("_", " ").title()

        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.add_run()
        run.add_picture(str(fig_path), width=Inches(5.5))

        cap = doc.add_paragraph(f"Figure {self.fig_counter}: {caption_txt}")
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.name   = FONT_NAME
            run.font.size   = Pt(FONT_SIZE_CAPTION)
            run.font.italic = True

        doc.add_paragraph()

    # ───────────────────────────────────────────
    # Utilities
    # ───────────────────────────────────────────

    def _add_page_break(self):
        self.doc.add_page_break()

    # ───────────────────────────────────────────
    # Excel export
    # ───────────────────────────────────────────

    def build_excel(self) -> Path:
        """Export all computed tables to a single Excel workbook."""
        from config import OUTPUT_EXCEL
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        path = OUTPUT_EXCEL / f"{self.p.get('project_name','Project').replace(' ','_')}_{self.doc_key}.xlsx"
        tg   = TableGenerator()

        tables = {
            "Project Data":          tg.project_data_table(self.model, self.cs),
            "Operational Params":    tg.operational_parameters_table(self.model, self.cs),
            "Rolling Stock":         tg.rolling_stock_table(self.model, self.cs),
            "Fleet Calculation":     tg.fleet_calculation_table(self.model, self.cs),
            "Headway Breakdown":     tg.headway_breakdown_table(self.cs),
            "RAM Targets":           tg.ram_targets_table(self.model, self.cs),
            "FMECA":                 tg.fmeca_table(),
            "Hazard Log":            tg.hazard_log_table(),
            "Interface Matrix":      tg.interface_matrix_table(),
            "SRS Requirements":      tg.srs_requirements_table(self.model, self.cs),
            "Standards":             tg.standards_table(self.model),
            "Env Conditions":        tg.environmental_conditions_table(self.model),
        }

        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            for sheet_name, df in tables.items():
                df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
                ws = writer.sheets[sheet_name[:31]]
                # Style header row
                header_fill = PatternFill("solid", fgColor="003087")
                for cell in ws[1]:
                    cell.font      = Font(bold=True, color="FFFFFF", name="Arial", size=10)
                    cell.fill      = header_fill
                    cell.alignment = Alignment(horizontal="center", wrap_text=True)
                # Auto-width
                for col in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col), default=0)
                    ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
                # Alternate row colours
                alt_fill = PatternFill("solid", fgColor="EAF0FB")
                for i, row in enumerate(ws.iter_rows(min_row=2), start=2):
                    if i % 2 == 0:
                        for cell in row:
                            cell.fill = alt_fill

        return path
