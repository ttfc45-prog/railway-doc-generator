"""
traceability_engine.py
Railway Documentation Generator — Traceability Engine
=====================================================

For every calculated parameter, records:
  • Parameter name and symbol
  • Authoritative source (ProjectModel input or CalculatedState field)
  • Formula (in engineering notation)
  • Calculation function (Python path)
  • Standard reference
  • Documents that use this parameter
  • Current value (from CalculatedState)

Generates:
  • HTML traceability report (self-contained)
  • Word .docx traceability matrix

DESIGN PRINCIPLE:
  Every engineering number in every document must appear in this matrix.
  If a value is used in a document but not in this matrix → architecture violation.
"""

from __future__ import annotations

import datetime
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


# ═══════════════════════════════════════════════════════════════════════════════
# TRACEABILITY ENTRY
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class TraceEntry:
    """Single parameter in the traceability matrix."""
    param_id:       str         # e.g. "OP-001"
    category:       str         # Operational / Headway / Capacity / RAM / Traction / RAMS
    symbol:         str         # engineering symbol (e.g. "v_c")
    name:           str         # full name
    unit:           str         # SI unit
    formula:        str         # engineering formula (LaTeX-style text)
    source:         str         # "ProjectModel.inputs" or "CalculationEngine"
    calc_function:  str         # Python function path
    standard_ref:   str         # EN/IEC standard reference
    documents:      list[str]   # documents that display this value
    value:          Any = None  # current calculated value (filled at runtime)
    value_str:      str = ""    # formatted string representation


# ═══════════════════════════════════════════════════════════════════════════════
# PARAMETER REGISTRY
# ═══════════════════════════════════════════════════════════════════════════════

PARAMETER_REGISTRY: list[TraceEntry] = [

    # ── OPERATIONAL ──────────────────────────────────────────────────────────

    TraceEntry(
        param_id      = "OP-001",
        category      = "Operational",
        symbol        = "v_c",
        name          = "Commercial Speed",
        unit          = "km/h",
        formula       = "v_c = L / t_running  [L in km, t_running in hours]",
        source        = "CalculationEngine._calculate_operations",
        calc_function = "calculations.CalculationEngine._calculate_operations",
        standard_ref  = "EN 62290-1 §3.1, UIC 406",
        documents     = ["ConOps", "OCS", "POP", "BOD", "SRS", "FleetCalc",
                         "HeadwayStudy", "CapacityStudy", "ExecutiveSummary"],
    ),
    TraceEntry(
        param_id      = "OP-002",
        category      = "Operational",
        symbol        = "t_run",
        name          = "Running Time (one way)",
        unit          = "min",
        formula       = "t_run = Σ(t_acc_i + t_cruise_i + t_dec_i) + N_stops × t_dwell",
        source        = "CalculationEngine._calculate_operations",
        calc_function = "calculations.CalculationEngine._calculate_operations",
        standard_ref  = "UIC 406 timetable compression",
        documents     = ["ConOps", "FleetCalc", "HeadwayStudy", "CapacityStudy"],
    ),
    TraceEntry(
        param_id      = "OP-003",
        category      = "Operational",
        symbol        = "RTT",
        name          = "Round Trip Time",
        unit          = "min",
        formula       = "RTT = 2 × t_run + 2 × t_terminal",
        source        = "CalculationEngine._calculate_operations",
        calc_function = "calculations.CalculationEngine._calculate_operations",
        standard_ref  = "EN 62290-1 §5.4",
        documents     = ["ConOps", "OCS", "FleetCalc", "HeadwayStudy"],
    ),
    TraceEntry(
        param_id      = "OP-004",
        category      = "Operational",
        symbol        = "N_s",
        name          = "Trains in Peak Service",
        unit          = "trains",
        formula       = "N_s = ⌈RTT_sec / H_sec⌉",
        source        = "CalculationEngine._calculate_operations",
        calc_function = "calculations.CalculationEngine._calculate_operations",
        standard_ref  = "EN 62290-1 §5.3",
        documents     = ["ConOps", "FleetCalc", "CapacityStudy"],
    ),
    TraceEntry(
        param_id      = "OP-005",
        category      = "Operational",
        symbol        = "N_fleet",
        name          = "Operational Fleet",
        unit          = "trains",
        formula       = "N_fleet = ⌈N_s / A_operational⌉",
        source        = "CalculationEngine._calculate_operations",
        calc_function = "calculations.CalculationEngine._calculate_operations",
        standard_ref  = "EN 62290-1 §5.3, EN 50126-1 §4",
        documents     = ["ConOps", "FleetCalc", "SRS", "ExecutiveSummary"],
    ),
    TraceEntry(
        param_id      = "OP-006",
        category      = "Operational",
        symbol        = "N_total",
        name          = "Total Fleet (incl. Reserve)",
        unit          = "trains",
        formula       = "N_total = N_fleet + N_reserve  where N_reserve = ⌈N_fleet × r_reserve⌉",
        source        = "CalculationEngine._calculate_operations",
        calc_function = "calculations.CalculationEngine._calculate_operations",
        standard_ref  = "EN 62290-1 §5.3",
        documents     = ["ConOps", "FleetCalc", "SRS", "BOD", "ExecutiveSummary"],
    ),
    TraceEntry(
        param_id      = "OP-007",
        category      = "Operational",
        symbol        = "TKM_yr",
        name          = "Annual Train-km",
        unit          = "train-km/year",
        formula       = "TKM_yr = N_s × (H_op/RTT) × 2L × 365",
        source        = "CalculationEngine._calculate_operations",
        calc_function = "calculations.CalculationEngine._calculate_operations",
        standard_ref  = "UIC 406",
        documents     = ["ConOps", "MaintenancePlan", "Performance", "RAM"],
    ),

    # ── HEADWAY ──────────────────────────────────────────────────────────────

    TraceEntry(
        param_id      = "HW-001",
        category      = "Headway",
        symbol        = "H_tech",
        name          = "Technical Headway",
        unit          = "s",
        formula       = "H_tech = t_react + t_tx + t_brk + t_margin + t_jerk  "
                        "where t_brk = v_max / d_emg",
        source        = "CalculationEngine._calculate_headway",
        calc_function = "calculations.CalculationEngine._calculate_headway",
        standard_ref  = "EN 62290-1 §5.4, IEC 62290-2",
        documents     = ["SRS", "HeadwayStudy", "SignallingDesc", "ConOps",
                         "FleetCalc", "CapacityStudy"],
    ),
    TraceEntry(
        param_id      = "HW-002",
        category      = "Headway",
        symbol        = "H_comm",
        name          = "Commercial Headway",
        unit          = "s",
        formula       = "H_comm = H_tech + t_dwell",
        source        = "CalculationEngine._calculate_headway",
        calc_function = "calculations.CalculationEngine._calculate_headway",
        standard_ref  = "EN 62290-1 §5.4",
        documents     = ["HeadwayStudy", "ConOps", "OCS"],
    ),
    TraceEntry(
        param_id      = "HW-003",
        category      = "Headway",
        symbol        = "d_sep",
        name          = "Minimum Safe Separation Distance",
        unit          = "m",
        formula       = "d_sep = v_max²/(2×d_emg) + v_max×(t_react + t_tx)",
        source        = "CalculationEngine._calculate_headway",
        calc_function = "calculations.CalculationEngine._calculate_headway",
        standard_ref  = "EN 62290-1 §5.4, EN 50159",
        documents     = ["HeadwayStudy", "SignallingDesc", "SRS"],
    ),

    # ── CAPACITY ─────────────────────────────────────────────────────────────

    TraceEntry(
        param_id      = "CA-001",
        category      = "Capacity",
        symbol        = "C_6",
        name          = "Train Capacity (6 pax/m²)",
        unit          = "passengers",
        formula       = "C_6 = seated + standing_6ppm2  [from ProjectModel]",
        source        = "ProjectModel.inputs",
        calc_function = "calculations.CalculationEngine._calculate_capacity",
        standard_ref  = "EN 15663, EN 13452",
        documents     = ["ConOps", "SRS", "CapacityStudy", "RollingStockDesc"],
    ),
    TraceEntry(
        param_id      = "CA-002",
        category      = "Capacity",
        symbol        = "PPHPD_6",
        name          = "Line Capacity at 6 pax/m²",
        unit          = "pphpd",
        formula       = "PPHPD_6 = C_6 × 3600 / H_service  [H_service in seconds]",
        source        = "CalculationEngine._calculate_capacity",
        calc_function = "calculations.CalculationEngine._calculate_capacity",
        standard_ref  = "EN 62290-1 §5.3, UITP Capacity Guide",
        documents     = ["ConOps", "SRS", "CapacityStudy", "BOD", "ExecutiveSummary"],
    ),
    TraceEntry(
        param_id      = "CA-003",
        category      = "Capacity",
        symbol        = "LF_6",
        name          = "Load Factor at 6 pax/m²",
        unit          = "%",
        formula       = "LF_6 = Peak_demand / PPHPD_6 × 100",
        source        = "CalculationEngine._calculate_capacity",
        calc_function = "calculations.CalculationEngine._calculate_capacity",
        standard_ref  = "UITP Statistics",
        documents     = ["CapacityStudy", "ConOps", "BOD"],
    ),

    # ── RAM ──────────────────────────────────────────────────────────────────

    TraceEntry(
        param_id      = "RA-001",
        category      = "RAM",
        symbol        = "A",
        name          = "System Availability",
        unit          = "%",
        formula       = "A = MTBF / (MTBF + MTTR)",
        source        = "CalculationEngine._calculate_ram",
        calc_function = "calculations.CalculationEngine._calculate_ram",
        standard_ref  = "EN 50126-1 §6.1",
        documents     = ["RAM", "SRS", "MaintenancePlan", "ExecutiveSummary",
                         "ConOps", "BOD"],
    ),
    TraceEntry(
        param_id      = "RA-002",
        category      = "RAM",
        symbol        = "R(24h)",
        name          = "Mission Reliability (24h)",
        unit          = "–",
        formula       = "R(t) = exp(−λ×t)  where λ = 1/MTBF, t = 24 h",
        source        = "CalculationEngine._calculate_ram",
        calc_function = "calculations.CalculationEngine._calculate_ram",
        standard_ref  = "EN 50126-1 §6.2",
        documents     = ["RAM", "Reliability"],
    ),
    TraceEntry(
        param_id      = "RA-003",
        category      = "RAM",
        symbol        = "M(8h)",
        name          = "Maintainability (8h shift reference)",
        unit          = "–",
        formula       = "M(t_ref) = 1 − exp(−μ×t_ref)  where μ = 1/MTTR, t_ref = 8 h",
        source        = "CalculationEngine._calculate_ram",
        calc_function = "calculations.CalculationEngine._calculate_ram",
        standard_ref  = "EN 50126-2 §6.3.3",
        documents     = ["RAM", "Maintainability", "MaintenancePlan"],
    ),
    TraceEntry(
        param_id      = "RA-004",
        category      = "RAM",
        symbol        = "MKBF",
        name          = "km Between Failures",
        unit          = "km",
        formula       = "MKBF = MTBF × v_c  [v_c from OperationalResults — correct chain]",
        source        = "CalculationEngine._calculate_ram",
        calc_function = "calculations.CalculationEngine._calculate_ram",
        standard_ref  = "EN 50126-1 §6.2, UITP Benchmarking",
        documents     = ["RAM", "SRS", "MaintenancePlan", "ConOps"],
    ),

    # ── TRACTION ─────────────────────────────────────────────────────────────

    TraceEntry(
        param_id      = "TR-001",
        category      = "Traction",
        symbol        = "P_peak",
        name          = "Peak Tractive Power per Train",
        unit          = "kW",
        formula       = "P = F_acc × v_max / η_drive  "
                        "where F_acc = m×ρ_rot×(a + g×r_davis)",
        source        = "CalculationEngine._calculate_traction",
        calc_function = "calculations.CalculationEngine._calculate_traction",
        standard_ref  = "EN 50163, IEC 61992",
        documents     = ["TractionDesc", "BOD", "SRS"],
    ),
    TraceEntry(
        param_id      = "TR-002",
        category      = "Traction",
        symbol        = "E_net",
        name          = "Net Energy per Train-km",
        unit          = "kWh/km",
        formula       = "E_net = (E_kin + E_res + E_grad + E_aux − E_regen)  "
                        "E_kin = ½mρ_rot v²×stops/km / η_drive; "
                        "E_res = F_davis×1000 / η_drive / 3.6e6; "
                        "E_grad = mg×grade×1000 / η_drive / 3.6e6; "
                        "E_aux = P_aux / v_c; "
                        "E_regen = E_kin×η_regen×f_regen",
        source        = "CalculationEngine._calculate_traction",
        calc_function = "calculations.CalculationEngine._calculate_traction",
        standard_ref  = "EN 50641, UIC 544-1, Davis (1926)",
        documents     = ["TractionDesc", "BOD", "Performance"],
    ),
    TraceEntry(
        param_id      = "TR-003",
        category      = "Traction",
        symbol        = "η_regen",
        name          = "Regenerative Energy Saving",
        unit          = "%",
        formula       = "saving = E_regen / E_gross × 100  "
                        "E_regen = E_kin × f_regen × η_regen",
        source        = "CalculationEngine._calculate_traction",
        calc_function = "calculations.CalculationEngine._calculate_traction",
        standard_ref  = "EN 50641 §6.3",
        documents     = ["TractionDesc", "BOD", "RollingStockDesc"],
    ),
    TraceEntry(
        param_id      = "TR-004",
        category      = "Traction",
        symbol        = "E_yr",
        name          = "Annual Energy Consumption",
        unit          = "MWh/year",
        formula       = "E_yr = TKM_yr × E_net / 1000  [uses ops.annual_train_km]",
        source        = "CalculationEngine._calculate_traction",
        calc_function = "calculations.CalculationEngine._calculate_traction",
        standard_ref  = "EN 50641",
        documents     = ["TractionDesc", "BOD", "Performance"],
    ),
    TraceEntry(
        param_id      = "TR-005",
        category      = "Traction",
        symbol        = "S_subs",
        name          = "Substation Rating",
        unit          = "MVA",
        formula       = "S = (N_fleet/N_subs) × P_peak × D_factor / PF / 1000  "
                        "D_factor=0.60, PF=0.90",
        source        = "CalculationEngine._calculate_traction",
        calc_function = "calculations.CalculationEngine._calculate_traction",
        standard_ref  = "EN 50163 §5.4, EN 50329",
        documents     = ["TractionDesc", "BOD"],
    ),

    # ── RAMS ALLOCATION ───────────────────────────────────────────────────────

    TraceEntry(
        param_id      = "RS-001",
        category      = "RAMS Allocation",
        symbol        = "MTBF_i",
        name          = "Subsystem Allocated MTBF",
        unit          = "h",
        formula       = "MTBF_i = MTBF_sys × Σw_j / w_i  "
                        "[Σ(1/MTBF_i) = 1/MTBF_sys — series constraint by construction]",
        source        = "CalculationEngine._calculate_rams_allocation",
        calc_function = "calculations.CalculationEngine._calculate_rams_allocation",
        standard_ref  = "EN 50126-2 §6.2 — RAMS apportionment",
        documents     = ["RAM", "Maintainability", "Reliability"],
    ),
    TraceEntry(
        param_id      = "RS-002",
        category      = "RAMS Allocation",
        symbol        = "A_i",
        name          = "Subsystem Allocated Availability",
        unit          = "%",
        formula       = "A_i = MTBF_i / (MTBF_i + MTTR_i)  "
                        "Verification: Π(A_i) ≥ A_sys_target",
        source        = "CalculationEngine._calculate_rams_allocation",
        calc_function = "calculations.CalculationEngine._calculate_rams_allocation",
        standard_ref  = "EN 50126-2 §6.2",
        documents     = ["RAM", "Availability"],
    ),
    TraceEntry(
        param_id      = "RS-003",
        category      = "RAMS Allocation",
        symbol        = "MTBF_series",
        name          = "Series Equivalent System MTBF",
        unit          = "h",
        formula       = "MTBF_series = 1/Σ(1/MTBF_i)  [should equal MTBF_sys target]",
        source        = "CalculationEngine._calculate_rams_allocation",
        calc_function = "calculations.CalculationEngine._calculate_rams_allocation",
        standard_ref  = "EN 50126-2 §6.2 Verification",
        documents     = ["RAM"],
    ),
]


# ═══════════════════════════════════════════════════════════════════════════════
# VALUE POPULATION
# ═══════════════════════════════════════════════════════════════════════════════

def populate_values(cs) -> list[TraceEntry]:
    """Populate each TraceEntry.value from the CalculatedState."""
    import copy
    entries = copy.deepcopy(PARAMETER_REGISTRY)

    value_map = {
        "OP-001": (cs.ops.commercial_speed_kmh,         f"{cs.ops.commercial_speed_kmh:.2f} km/h"),
        "OP-002": (cs.ops.running_time_min,              f"{cs.ops.running_time_min:.1f} min"),
        "OP-003": (cs.ops.round_trip_time_min,           f"{cs.ops.round_trip_time_min:.1f} min"),
        "OP-004": (cs.ops.trains_in_service,             str(cs.ops.trains_in_service)),
        "OP-005": (cs.ops.fleet_required,                str(cs.ops.fleet_required)),
        "OP-006": (cs.ops.total_fleet,                   str(cs.ops.total_fleet)),
        "OP-007": (cs.ops.annual_train_km,               f"{cs.ops.annual_train_km:,.0f}"),
        "HW-001": (cs.headway.technical_headway_sec,     f"{cs.headway.technical_headway_sec:.1f} s"),
        "HW-002": (cs.headway.commercial_headway_sec,    f"{cs.headway.commercial_headway_sec:.1f} s"),
        "HW-003": (cs.headway.minimum_safe_separation_m, f"{cs.headway.minimum_safe_separation_m:.0f} m"),
        "CA-001": (cs.capacity.capacity_6ppm2,           f"{cs.capacity.capacity_6ppm2} pax"),
        "CA-002": (cs.capacity.pphpd_6ppm2,              f"{cs.capacity.pphpd_6ppm2:,} pphpd"),
        "CA-003": (cs.capacity.load_factor_6ppm2_pct,    f"{cs.capacity.load_factor_6ppm2_pct:.1f}%"),
        "RA-001": (cs.ram.availability * 100,            f"{cs.ram.availability*100:.4f}%"),
        "RA-002": (cs.ram.mission_reliability_24h,       f"{cs.ram.mission_reliability_24h:.5f}"),
        "RA-003": (cs.ram.maintainability_8h,            f"{cs.ram.maintainability_8h:.4f}"),
        "RA-004": (cs.ram.km_between_failures,           f"{cs.ram.km_between_failures:,.0f} km"),
        "TR-001": (cs.traction.peak_power_kw,            f"{cs.traction.peak_power_kw:,.0f} kW"),
        "TR-002": (cs.traction.energy_per_train_km_kwh,  f"{cs.traction.energy_per_train_km_kwh:.3f} kWh/km"),
        "TR-003": (cs.traction.regenerative_saving_pct,  f"{cs.traction.regenerative_saving_pct:.1f}%"),
        "TR-004": (cs.traction.annual_energy_mwh,        f"{cs.traction.annual_energy_mwh:,.0f} MWh/yr"),
        "TR-005": (cs.traction.substation_rating_mva,    f"{cs.traction.substation_rating_mva:.1f} MVA"),
        "RS-001": (None,                                  "See subsystem allocation table"),
        "RS-002": (None,                                  "See subsystem allocation table"),
        "RS-003": (cs.rams_alloc.series_mtbf_hours,      f"{cs.rams_alloc.series_mtbf_hours:,.0f} h"),
    }

    for e in entries:
        if e.param_id in value_map:
            e.value, e.value_str = value_map[e.param_id]

    return entries


# ═══════════════════════════════════════════════════════════════════════════════
# HTML REPORT RENDERER
# ═══════════════════════════════════════════════════════════════════════════════

def render_html_traceability(model, cs) -> str:
    """Render a self-contained HTML traceability report."""
    entries = populate_values(cs)
    now = datetime.datetime.now().strftime("%d %B %Y  %H:%M")
    p   = model.to_dict()

    # Group by category
    from itertools import groupby
    sorted_entries = sorted(entries, key=lambda e: e.category)

    cat_colours = {
        "Operational":    "#003087",
        "Headway":        "#1A5276",
        "Capacity":       "#154360",
        "RAM":            "#7B241C",
        "Traction":       "#145A32",
        "RAMS Allocation":"#4A235A",
    }

    rows_html = ""
    current_cat = None
    for e in sorted_entries:
        if e.category != current_cat:
            current_cat = e.category
            bg = cat_colours.get(current_cat, "#333")
            rows_html += (
                f"<tr><td colspan='8' style='background:{bg};color:#fff;"
                f"padding:8px 14px;font-size:11px;font-weight:800;"
                f"letter-spacing:.1em;text-transform:uppercase'>"
                f"▸ {current_cat}</td></tr>\n"
            )

        docs_str = "<br>".join(e.documents[:5])
        if len(e.documents) > 5:
            docs_str += f"<br><em>+{len(e.documents)-5} more</em>"

        rows_html += f"""
        <tr>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              font-weight:700;color:#003087;white-space:nowrap'>{e.param_id}</td>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              font-family:monospace;font-size:13px;color:#1A5276'>{e.symbol}</td>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              font-weight:600'>{e.name}</td>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              color:#374151;font-size:11px;max-width:280px;word-break:break-word'>{e.formula}</td>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              font-family:monospace;font-size:11px;color:#6B7280'>{e.calc_function}</td>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              font-size:11px;color:#6B7280'>{e.standard_ref}</td>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              font-size:11px'>{docs_str}</td>
          <td style='padding:8px 10px;border-bottom:1px solid #E5E7EB;
              font-weight:800;color:#155724;white-space:nowrap'>{e.value_str}</td>
        </tr>\n"""

    # RAMS allocation sub-table
    ra = cs.rams_alloc
    rams_rows = ""
    for name, mtbf, mttr, avail, w in zip(
        ra.subsystem_names, ra.allocated_mtbf_hours,
        ra.allocated_mttr_hours, ra.allocated_avail_pct, ra.complexity_weights
    ):
        rams_rows += (
            f"<tr>"
            f"<td style='padding:6px 10px;border-bottom:1px solid #E5E7EB'>{name}</td>"
            f"<td style='padding:6px 10px;border-bottom:1px solid #E5E7EB;text-align:center'>{w:.1f}</td>"
            f"<td style='padding:6px 10px;border-bottom:1px solid #E5E7EB;font-family:monospace'>{int(mtbf):,} h</td>"
            f"<td style='padding:6px 10px;border-bottom:1px solid #E5E7EB;font-family:monospace'>{mttr:.1f} h</td>"
            f"<td style='padding:6px 10px;border-bottom:1px solid #E5E7EB;font-family:monospace'>{avail:.4f}%</td>"
            f"</tr>\n"
        )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>Traceability Matrix — {p.get('project_name','')}</title>
<style>
  body {{font-family:"Segoe UI",Arial,sans-serif;margin:0;background:#fff;color:#1B2631;font-size:13px}}
  table {{border-collapse:collapse;width:100%}}
  th {{background:#003087;color:#fff;padding:9px 10px;text-align:left;
       font-size:11px;letter-spacing:.06em;text-transform:uppercase}}
  .hdr {{background:#003087;color:#fff;padding:28px 40px;border-bottom:5px solid #C0392B}}
  .hdr h1 {{font-size:20px;font-weight:800;margin:0 0 6px;letter-spacing:-.01em}}
  .sec {{padding:20px 30px}}
  .sec h2 {{font-size:14px;font-weight:800;color:#003087;border-bottom:2px solid #003087;
            padding-bottom:5px;margin-bottom:12px}}
  .badge {{display:inline-block;padding:2px 8px;border-radius:3px;font-size:10px;
           font-weight:800;letter-spacing:.08em;text-transform:uppercase}}
  .pass {{background:#D5F5E3;color:#145A32;border:1px solid #82E0AA}}
  .warn {{background:#FEF9E7;color:#7D6608;border:1px solid #F9E79F}}
</style>
</head>
<body>
<div class="hdr">
  <div style="font-size:10px;letter-spacing:.18em;text-transform:uppercase;opacity:.6;margin-bottom:4px">
    Engineering Traceability Matrix · Phase 1.5
  </div>
  <h1>Parameter Traceability Matrix</h1>
  <div style="font-size:12px;opacity:.7;margin-top:6px">
    Project: {p.get('project_name','')} &nbsp;|&nbsp;
    Generated: {now} &nbsp;|&nbsp;
    Parameters: {len(entries)} &nbsp;|&nbsp;
    Source: CalculationEngine (deterministic)
  </div>
</div>

<div class="sec">
  <h2>Engineering Summary</h2>
  <table style="width:auto;margin-bottom:20px">
    <tr><td style="padding:5px 14px;font-weight:700">Commercial Speed</td>
        <td style="padding:5px 14px;font-family:monospace">{cs.ops.commercial_speed_kmh:.2f} km/h</td></tr>
    <tr style="background:#F8F9FA"><td style="padding:5px 14px;font-weight:700">Technical Headway</td>
        <td style="padding:5px 14px;font-family:monospace">{cs.headway.technical_headway_sec:.1f} s</td></tr>
    <tr><td style="padding:5px 14px;font-weight:700">Line Capacity (6 pax/m²)</td>
        <td style="padding:5px 14px;font-family:monospace">{cs.capacity.pphpd_6ppm2:,} pphpd 
        <span class="badge {'pass' if cs.capacity.capacity_adequate else 'warn'}">
        {'≥ demand ✓' if cs.capacity.capacity_adequate else '< demand ✗'}</span></td></tr>
    <tr style="background:#F8F9FA"><td style="padding:5px 14px;font-weight:700">Net Energy / Train-km</td>
        <td style="padding:5px 14px;font-family:monospace">{cs.traction.energy_per_train_km_kwh:.3f} kWh/km (full physics)</td></tr>
    <tr><td style="padding:5px 14px;font-weight:700">Series MTBF (subsystems)</td>
        <td style="padding:5px 14px;font-family:monospace">{cs.rams_alloc.series_mtbf_hours:,.0f} h
        <span class="badge {'pass' if cs.rams_alloc.series_mtbf_consistent else 'warn'}">
        {'= target ✓' if cs.rams_alloc.series_mtbf_consistent else '≠ target'}</span></td></tr>
    <tr style="background:#F8F9FA"><td style="padding:5px 14px;font-weight:700">Series Availability (subsystems)</td>
        <td style="padding:5px 14px;font-family:monospace">{cs.rams_alloc.series_avail_pct:.4f}%
        <span class="badge {'pass' if cs.rams_alloc.series_meets_avail_target else 'warn'}">
        {'≥ target ✓' if cs.rams_alloc.series_meets_avail_target else '< target'}</span></td></tr>
  </table>
</div>

<div class="sec">
  <h2>Complete Parameter Traceability Matrix ({len(entries)} parameters)</h2>
  <table>
    <thead>
      <tr>
        <th style="width:70px">ID</th>
        <th style="width:50px">Symbol</th>
        <th style="width:160px">Parameter</th>
        <th style="width:300px">Formula</th>
        <th style="width:180px">Calc Function</th>
        <th style="width:130px">Standard Ref.</th>
        <th style="width:130px">Documents</th>
        <th style="width:100px">Current Value</th>
      </tr>
    </thead>
    <tbody>{rows_html}</tbody>
  </table>
</div>

<div class="sec">
  <h2>RAMS Subsystem Allocation Detail</h2>
  <p style="font-size:12px;color:#555;margin-bottom:10px">
    Method: MTBF_i = MTBF_sys × Σw_j / w_i &nbsp;|&nbsp;
    Series MTBF = {cs.rams_alloc.series_mtbf_hours:,.0f} h (target: {cs.rams_alloc.system_mtbf_target:,.0f} h) &nbsp;|&nbsp;
    Series A = {cs.rams_alloc.series_avail_pct:.4f}% (target: {cs.rams_alloc.system_avail_target_pct}%) &nbsp;|&nbsp;
    Standard: EN 50126-2 §6.2
  </p>
  <table>
    <thead>
      <tr>
        <th>Subsystem</th>
        <th>Weight</th>
        <th>Allocated MTBF</th>
        <th>MTTR</th>
        <th>Availability</th>
      </tr>
    </thead>
    <tbody>{rams_rows}</tbody>
  </table>
</div>

<div class="sec" style="color:#888;font-size:11px;border-top:1px solid #E5E7EB;padding-top:12px">
  Railway Documentation Generator — Traceability Engine · Phase 1.5 ·
  All values from CalculationEngine (deterministic, frozen CalculatedState).
</div>
</body>
</html>"""


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX TRACEABILITY REPORT
# ═══════════════════════════════════════════════════════════════════════════════

def render_docx_traceability(model, cs) -> Path:
    """Generate a Word .docx traceability matrix document."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from config import OUTPUT_DOCX

    entries = populate_values(cs)
    p       = model.to_dict()

    def _rgb(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))
    def _cell_bg(cell, h):
        tc = cell._tc; pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), h)
        pr.append(shd)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)   # A3 landscape
    for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
        setattr(sec, attr, Cm(1.8))

    # Title
    t = doc.add_heading("Parameter Traceability Matrix", level=1)
    t.runs[0].font.color.rgb = _rgb("003087")
    doc.add_paragraph(
        f"Project: {p.get('project_name','')}  |  "
        f"Generated: {datetime.datetime.now().strftime('%d %B %Y')}  |  "
        f"Parameters: {len(entries)}  |  Source: CalculationEngine"
    ).runs[0].font.size = Pt(9)

    doc.add_heading("Traceability Matrix", level=2)

    cols = ["ID","Symbol","Parameter","Formula","Calc Function",
            "Standard","Documents","Value"]
    widths_cm = [1.5, 1.5, 3.5, 6.5, 4.0, 3.0, 3.0, 2.5]

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    hrow = table.rows[0]
    for i, (h, w) in enumerate(zip(cols, widths_cm)):
        cell = hrow.cells[i]
        _cell_bg(cell, "003087")
        run = cell.paragraphs[0].add_run(h)
        run.font.color.rgb = _rgb("FFFFFF"); run.font.bold = True
        run.font.size = Pt(8)
        cell.width = Cm(w)

    cat_colours_docx = {
        "Operational":"003087", "Headway":"1A5276", "Capacity":"154360",
        "RAM":"7B241C",         "Traction":"145A32", "RAMS Allocation":"4A235A",
    }
    sorted_entries = sorted(entries, key=lambda e: e.category)
    current_cat = None

    for e in sorted_entries:
        if e.category != current_cat:
            current_cat = e.category
            crow = table.add_row()
            merged = crow.cells[0].merge(crow.cells[-1])
            _cell_bg(merged, cat_colours_docx.get(current_cat,"333333"))
            run = merged.paragraphs[0].add_run(f"▸ {current_cat.upper()}")
            run.font.color.rgb = _rgb("FFFFFF"); run.font.bold = True
            run.font.size = Pt(8)

        row = table.add_row()
        vals = [
            e.param_id, e.symbol, e.name, e.formula[:80],
            e.calc_function.split(".")[-1],
            e.standard_ref[:30], ", ".join(e.documents[:3]),
            e.value_str
        ]
        for i, (v, w) in enumerate(zip(vals, widths_cm)):
            cell = row.cells[i]; cell.width = Cm(w)
            run = cell.paragraphs[0].add_run(str(v))
            run.font.size = Pt(7.5)
            if i == 0: run.font.bold = True; run.font.color.rgb = _rgb("003087")
            if i == 7: run.font.bold = True; run.font.color.rgb = _rgb("155724")

    path = OUTPUT_DOCX / f"{p.get('project_name','Project').replace(' ','_')}_TraceabilityMatrix.docx"
    doc.save(str(path))
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════════════════

def render_traceability_tab(model, cs) -> None:
    """Render the traceability matrix inside a Streamlit tab."""
    import streamlit as st
    import pandas as pd

    st.markdown("### Parameter Traceability Matrix")
    st.caption(
        "Every engineering number in every document traced to its formula, "
        "calculation function, and standard reference."
    )

    entries = populate_values(cs)

    # Summary KPIs
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Parameters", len(entries))
    c2.metric("Series MTBF", f"{cs.rams_alloc.series_mtbf_hours:,.0f} h",
              help="Must equal MTBF target by construction")
    c3.metric("Series Avail.", f"{cs.rams_alloc.series_avail_pct:.4f}%",
              help="Product of subsystem availabilities (conservative lower bound)")
    c4.metric("Net Energy", f"{cs.traction.energy_per_train_km_kwh:.3f} kWh/km",
              help="Full physics model (kinetic + resistance + gradient + auxiliary − regen)")

    # Full matrix as dataframe
    rows = []
    for e in sorted(entries, key=lambda x: x.category):
        rows.append({
            "ID":         e.param_id,
            "Category":   e.category,
            "Symbol":     e.symbol,
            "Parameter":  e.name,
            "Unit":       e.unit,
            "Value":      e.value_str,
            "Formula":    e.formula[:70] + "…" if len(e.formula) > 70 else e.formula,
            "Standard":   e.standard_ref,
            "Documents":  ", ".join(e.documents[:4]),
        })

    df = pd.DataFrame(rows)
    st.dataframe(df, use_container_width=True, hide_index=True, height=500)

    # RAMS allocation detail
    with st.expander("📐 RAMS Subsystem Allocation (EN 50126-2 §6.2)"):
        ra = cs.rams_alloc
        rams_rows = []
        for name, mtbf, mttr, avail, w in zip(
            ra.subsystem_names, ra.allocated_mtbf_hours,
            ra.allocated_mttr_hours, ra.allocated_avail_pct, ra.complexity_weights
        ):
            rams_rows.append({
                "Subsystem":         name,
                "Weight":            w,
                "Allocated MTBF (h)":f"{int(mtbf):,}",
                "MTTR (h)":          f"{mttr:.1f}",
                "Availability (%)":  f"{avail:.4f}",
            })
        rams_df = pd.DataFrame(rams_rows)
        st.dataframe(rams_df, use_container_width=True, hide_index=True)

        col1, col2 = st.columns(2)
        col1.metric("Series MTBF",
                    f"{ra.series_mtbf_hours:,.0f} h",
                    delta=f"Target: {ra.system_mtbf_target:,.0f} h")
        col2.metric("Series Availability",
                    f"{ra.series_avail_pct:.4f}%",
                    delta=f"Target: {ra.system_avail_target_pct}%")

        if ra.series_mtbf_consistent:
            st.success("✅ Series MTBF = system target (consistent by construction)")
        if ra.series_meets_avail_target:
            st.success("✅ Series availability ≥ system target")

    # Traction energy breakdown
    with st.expander("⚡ Traction Energy Model Breakdown (EN 50641)"):
        t = cs.traction
        breakdown = {
            "Component":       ["Kinetic Energy", "Rolling Resistance", "Gradient", "Auxiliary Loads", "Regen Recovery", "Net Energy"],
            "kWh/km":          [t.acc_energy_kwh_km, t.resistance_energy_kwh_km,
                                t.gradient_energy_kwh_km, t.auxiliary_energy_kwh_km,
                                -t.braking_energy_kwh_km, t.energy_per_train_km_kwh],
            "Formula":         ["½mρ_rot v²·stops/km / η",
                                "F_davis·1000 / η / 3.6e6",
                                "mg·grade·1000 / η / 3.6e6",
                                "P_aux / v_c",
                                "E_kin·f_regen·η_regen",
                                "Gross - E_regen"],
        }
        st.dataframe(pd.DataFrame(breakdown), use_container_width=True, hide_index=True)
        st.caption(f"Motor efficiency η = {t.motor_efficiency:.4f} (PMSM + gearbox) · "
                   f"Regen efficiency = {t.regen_efficiency:.2f}")

    # Export
    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        html_content = render_html_traceability(model, cs)
        st.download_button(
            "⬇  HTML Traceability Report",
            data=html_content,
            file_name=f"{model.get('project_name','Project').replace(' ','_')}_TraceabilityMatrix.html",
            mime="text/html",
            use_container_width=True,
        )
    with col2:
        if st.button("⬇  Generate DOCX Traceability Matrix", use_container_width=True):
            with st.spinner("Building Word document…"):
                docx_path = render_docx_traceability(model, cs)
                with open(docx_path, "rb") as f:
                    st.download_button(
                        "⬇  Download .docx",
                        data=f.read(),
                        file_name=docx_path.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
